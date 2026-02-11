"""
Generuje podręcznik użytkownika BABOK ANALYST w formacie DOCX.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ==================== STYLE SETUP ====================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = 'Calibri'
    h_style.font.color.rgb = RGBColor(0x1A, 0x47, 0x3C)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)

# Helper functions
def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return table

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run('💡 ')
    run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(1.27)
    return p

# ==================== TITLE PAGE ====================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('BABOK ANALYST')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x3C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Podręcznik Użytkownika')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run('Wersja 1.8.2')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('Luty 2026')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Narzędzie do profesjonalnej analizy biznesowej\nzgodne ze standardem BABOK® v3')
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ==================== TABLE OF CONTENTS ====================
doc.add_heading('Spis Treści', level=1)

toc_items = [
    '1. Wprowadzenie',
    '   1.1 Czym jest BABOK ANALYST?',
    '   1.2 Dla kogo jest to narzędzie?',
    '   1.3 Kluczowe cechy i możliwości',
    '   1.4 Framework BABOK® v3',
    '2. Wymagania i Przygotowanie',
    '   2.1 Wymagania systemowe',
    '   2.2 Konfiguracja agenta AI',
    '   2.3 Checklist przed rozpoczęciem',
    '3. Szybki Start — Pierwsze Kroki',
    '   3.1 Tworzenie nowego projektu',
    '   3.2 Zapisywanie i wznawianie projektu',
    '   3.3 Jak odpowiadać na pytania agenta',
    '   3.4 Zatwierdzanie i odrzucanie etapów',
    '4. Proces Analizy — 8 Etapów',
    '   4.1 Etap 1: Inicjalizacja Projektu i Mapa Interesariuszy',
    '   4.2 Etap 2: Analiza Stanu Obecnego (AS-IS)',
    '   4.3 Etap 3: Analiza Domeny Problemowej',
    '   4.4 Etap 4: Definicja Wymagań Rozwiązania',
    '   4.5 Etap 5: Projekt Stanu Docelowego (TO-BE)',
    '   4.6 Etap 6: Analiza Luk i Harmonogram Wdrożenia',
    '   4.7 Etap 7: Ocena Ryzyka i Strategia Mitygacji',
    '   4.8 Etap 8: Biznesplan i Model ROI',
    '5. Interfejs Poleceń — Pełna Referencja',
    '   5.1 Polecenia zarządzania projektem',
    '   5.2 Polecenia zarządzania etapami',
    '   5.3 Polecenia dokumentów',
    '   5.4 Polecenia analizy',
    '   5.5 Polecenia współpracy',
    '6. Struktura Projektu i Plików',
    '   6.1 Zalecana struktura folderów',
    '   6.2 Konwencja nazewnictwa plików',
    '   6.3 Wersjonowanie dokumentów',
    '   6.4 Bezpieczeństwo i kontrola dostępu',
    '7. Tryby Pracy Agenta',
    '   7.1 Tryb głębokiej analizy (Deep Analysis)',
    '   7.2 Tryb standardowy',
    '   7.3 Tryb szybki (Rapid Mode)',
    '8. Dziennik Projektu (Project Journal)',
    '9. Najlepsze Praktyki i Wskazówki',
    '10. Rozwiązywanie Problemów (FAQ)',
    '11. Bezpieczeństwo i Prywatność',
    '12. Szacunkowy Czas Realizacji',
    '13. Słowniczek Pojęć',
]

for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if not item.startswith('   '):
        for r in p.runs:
            r.font.bold = True

doc.add_page_break()

# ==================== CHAPTER 1: INTRODUCTION ====================
doc.add_heading('1. Wprowadzenie', level=1)

doc.add_heading('1.1 Czym jest BABOK ANALYST?', level=2)
doc.add_paragraph(
    'BABOK ANALYST to zaawansowane narzędzie wspomagające analizę biznesową dla projektów IT, '
    'oparte na międzynarodowym standardzie BABOK® v3 (Business Analysis Body of Knowledge) '
    'opracowanym przez International Institute of Business Analysis (IIBA). '
    'Narzędzie działa jako inteligentny agent AI, który prowadzi użytkownika przez kompleksowy, '
    '8-etapowy proces analizy biznesowej — od inicjalizacji projektu, przez mapowanie procesów, '
    'definicję wymagań, aż po biznesplan z modelem zwrotu z inwestycji (ROI).'
)
doc.add_paragraph(
    'Agent wykorzystuje model współpracy „human-in-the-loop" — oznacza to, że żaden etap nie jest '
    'finalizowany bez wyraźnego zatwierdzenia przez człowieka. Agent zadaje pytania sekwencyjnie, '
    'analizuje odpowiedzi, generuje profesjonalną dokumentację i oczekuje walidacji na każdym kroku.'
)

doc.add_heading('1.2 Dla kogo jest to narzędzie?', level=2)
doc.add_paragraph('BABOK ANALYST został zaprojektowany dla:')
add_bullet('Analityków biznesowych prowadzących projekty IT w firmach średniej wielkości')
add_bullet('Kierowników projektów odpowiedzialnych za wdrożenia systemów informatycznych')
add_bullet('Kadry zarządzającej (CEO, CFO) potrzebującej uzasadnienia biznesowego inwestycji IT')
add_bullet('Konsultantów IT pracujących z firmami z sektora produkcji, dystrybucji i usług')
add_bullet('Firm o przychodach €10–100M i zatrudnieniu 50–500 pracowników')

doc.add_heading('1.3 Kluczowe cechy i możliwości', level=2)

doc.add_paragraph('Specjalizacja branżowa:', style='List Bullet')
add_bullet('Produkcja, dystrybucja, usługi — sektor mid-market', level=1)
add_bullet('Regulacje UE/międzynarodowe: GDPR, e-fakturowanie, sprawozdawczość finansowa', level=1)

doc.add_paragraph('Inteligentny proces analizy:', style='List Bullet')
add_bullet('8 etapów zgodnych z BABOK® v3', level=1)
add_bullet('Adaptacyjna głębokość analizy — automatyczny wybór trybu AI', level=1)
add_bullet('Pytania zadawane sekwencyjnie z potwierdzeniem każdej odpowiedzi', level=1)

doc.add_paragraph('Profesjonalna dokumentacja:', style='List Bullet')
add_bullet('Automatyczne generowanie dokumentów Markdown na każdym etapie', level=1)
add_bullet('Eksport do DOCX i PDF za pomocą CLI', level=1)
add_bullet('Matryca RACI, rejestr ryzyk, RTM (Requirements Traceability Matrix)', level=1)

doc.add_paragraph('Zarządzanie projektem:', style='List Bullet')
add_bullet('Unikalny identyfikator projektu (BABOK-YYYYMMDD-XXXX)', level=1)
add_bullet('Dziennik projektu (Journal) z pełną historią zmian', level=1)
add_bullet('Zapis i wznowienie projektu w dowolnym momencie', level=1)

doc.add_heading('1.4 Framework BABOK® v3', level=2)
doc.add_paragraph(
    'BABOK® (Business Analysis Body of Knowledge) to międzynarodowy standard analizy biznesowej '
    'opracowany przez IIBA. Definiuje zbiór praktyk, technik i kompetencji wymaganych do skutecznej '
    'analizy biznesowej. BABOK ANALYST adaptuje ten framework do kontekstu firm mid-market, '
    'upraszczając formalizmy przy zachowaniu rygoru analitycznego.'
)
doc.add_paragraph('Wykorzystywane obszary wiedzy BABOK® v3:')
add_bullet('Business Analysis Planning and Monitoring — planowanie i kontrola analizy')
add_bullet('Elicitation and Collaboration — pozyskiwanie i walidacja informacji')
add_bullet('Requirements Life Cycle Management — zarządzanie cyklem życia wymagań')
add_bullet('Strategy Analysis — analiza strategiczna')
add_bullet('Requirements Analysis and Design Definition — analiza wymagań i definicja rozwiązań')
add_bullet('Solution Evaluation — ocena rozwiązań')

doc.add_page_break()

# ==================== CHAPTER 2: REQUIREMENTS ====================
doc.add_heading('2. Wymagania i Przygotowanie', level=1)

doc.add_heading('2.1 Wymagania systemowe', level=2)
doc.add_paragraph('Aby korzystać z BABOK ANALYST, potrzebujesz:')
add_bullet('Dostęp do platformy AI obsługującej system prompt (np. Claude.ai, API Anthropic, Gemini)')
add_bullet('Przeglądarka internetowa lub narzędzie API')
add_bullet('Opcjonalnie: Node.js (v18+) do korzystania z narzędzia wiersza poleceń (CLI)')
add_bullet('Opcjonalnie: VS Code z GitHub Copilot do pracy w środowisku IDE')

doc.add_heading('2.2 Konfiguracja agenta AI', level=2)
doc.add_paragraph('Krok 1: Przygotowanie system promptu')
doc.add_paragraph(
    'Skopiuj pełną zawartość pliku BABOK_Agent_System_Prompt.md i wklej ją do instrukcji projektu '
    '(Project Instructions / System Prompt) w wybranej platformie AI. Plik ten zawiera wszystkie '
    'instrukcje, szablony i zachowania agenta.'
)
doc.add_paragraph('Krok 2: Rozpoczęcie sesji')
doc.add_paragraph(
    'W nowej konwersacji wpisz polecenie BEGIN NEW PROJECT. Agent wygeneruje unikalny '
    'identyfikator projektu i rozpocznie etap 1.'
)
doc.add_paragraph('Krok 3: Opcjonalna konfiguracja CLI')
doc.add_paragraph(
    'Jeśli chcesz korzystać z narzędzia wiersza poleceń do zarządzania projektami i eksportu '
    'dokumentów, zainstaluj CLI wchodząc do katalogu cli/ i uruchamiając npm install.'
)

doc.add_heading('2.3 Checklist przed rozpoczęciem', level=2)
doc.add_paragraph('Przed rozpoczęciem analizy upewnij się, że:')
add_bullet('Masz 30–45 minut nieprzerwanego czasu na etap 1')
add_bullet('Posiadasz podstawowe informacje o firmie (przychody, liczba pracowników, systemy IT)')
add_bullet('Znasz sponsora projektu (kto podejmuje decyzje budżetowe)')
add_bullet('Masz możliwość konsultacji z zespołami (Finanse, IT, Operacje)')
add_bullet('Załadowałeś system prompt agenta BABOK do swojego narzędzia AI')
add_bullet('Przeczytałeś niniejszy podręcznik i rozumiesz proces')

doc.add_page_break()

# ==================== CHAPTER 3: QUICK START ====================
doc.add_heading('3. Szybki Start — Pierwsze Kroki', level=1)

doc.add_heading('3.1 Tworzenie nowego projektu', level=2)
doc.add_paragraph('Aby rozpocząć nowy projekt analizy biznesowej, wpisz polecenie:')
add_code_block('BEGIN NEW PROJECT')

doc.add_paragraph('Agent wykona następujące działania:')
add_bullet('Wygeneruje unikalny identyfikator projektu w formacie BABOK-YYYYMMDD-XXXX')
add_bullet('Utworzy katalog projektu i plik dziennika (journal)')
add_bullet('Wyświetli potwierdzenie z danymi projektu')
add_bullet('Przejdzie do etapu 1: Inicjalizacja Projektu')

doc.add_paragraph('Przykładowy wynik:')
add_code_block(
    '✅ NEW PROJECT CREATED\n'
    '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
    'Project ID:  BABOK-20260208-M3R1\n'
    'Created:     2026-02-08 10:30:00 UTC\n'
    '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
    'Proceeding to Stage 1...'
)

add_note('Zanotuj swój Project ID — będzie potrzebny do wznowienia projektu w przyszłości.')

doc.add_heading('3.2 Zapisywanie i wznawianie projektu', level=2)

doc.add_paragraph('Zapisywanie projektu (po zatwierdzeniu dowolnego etapu):')
add_code_block('SAVE PROJECT')

doc.add_paragraph('Wznowienie zapisanego projektu:')
add_code_block('LOAD PROJECT BABOK-20260208-M3R1')

doc.add_paragraph(
    'Agent odczyta dziennik projektu, przywróci pełen kontekst (dane ze wszystkich '
    'zatwierdzonych etapów, decyzje, założenia) i wznowi pracę dokładnie w miejscu, '
    'w którym została przerwana.'
)

doc.add_heading('3.3 Jak odpowiadać na pytania agenta', level=2)

doc.add_paragraph(
    'Agent zadaje pytania sekwencyjnie, po jednym na raz, z wyraźnym oznaczeniem numeru '
    'pytania (np. „Question 1/5"). Po każdej odpowiedzi agent krótko ją podsumuje '
    'i przejdzie do kolejnego pytania.'
)

doc.add_paragraph('Format pytań agenta:')
add_code_block(
    '📋 STAGE 1 - QUESTION 1/5\n\n'
    'Category: Project Scope - Document Types\n\n'
    'What types of documents are IN SCOPE for this project?\n\n'
    'Please provide your answer. I will proceed to question 2/5 once you respond.'
)

doc.add_paragraph('Zalecany sposób odpowiadania — jasno i konkretnie:')
add_code_block(
    '1. TAK - faktury zakupowe i sprzedażowe\n'
    '2. ERP: SAP Business One v10.0\n'
    '3. Nie mamy DMS obecnie\n'
    '4. KSeF deadline: 1 lipca 2026'
)

doc.add_paragraph('Jeśli czegoś nie wiesz:')
add_code_block('NIE WIEM - muszę sprawdzić z działem finansów')

doc.add_paragraph('Jeśli chcesz przyspieszyć (tryb batch):')
add_code_block('Skip questions')
doc.add_paragraph(
    'Polecenie „Skip questions" wyświetli wszystkie pozostałe pytania naraz, '
    'dzięki czemu możesz odpowiedzieć na nie hurtem.'
)

doc.add_heading('3.4 Zatwierdzanie i odrzucanie etapów', level=2)

doc.add_paragraph(
    'Po zebraniu wszystkich odpowiedzi agent generuje kompletny dokument (deliverable) '
    'danego etapu. Przed przejściem dalej musisz go zatwierdzić lub odrzucić.'
)

add_table(
    ['Akcja', 'Polecenie', 'Opis'],
    [
        ['Zatwierdzenie', 'Approve 1', 'Zatwierdza etap 1 i przechodzi do etapu 2'],
        ['Odrzucenie', 'Reject 1 brak danych bazowych', 'Odrzuca z podaniem powodu — agent poprawi'],
        ['Podgląd podsumowania', 'Summary 1', 'Wyświetla 1-stronicowe streszczenie wykonawcze'],
        ['Podgląd szczegółów', 'Detail 1', 'Wyświetla pełną analizę'],
        ['Walidacja kompletności', 'Validate 1', 'Sprawdza czy nic nie brakuje przed zatwierdzeniem'],
    ]
)

doc.add_page_break()

# ==================== CHAPTER 4: 8 STAGES ====================
doc.add_heading('4. Proces Analizy — 8 Etapów', level=1)

doc.add_paragraph(
    'Proces BABOK ANALYST składa się z 8 kolejnych etapów. Każdy etap wymaga zatwierdzenia '
    'przez człowieka (Human Approval) przed przejściem do następnego. Etapy 3, 4, 6 i 8 '
    'automatycznie aktywują tryb głębokiej analizy (Deep Analysis Mode), który wykorzystuje '
    'najbardziej zaawansowane modele AI do krytycznych decyzji.'
)

doc.add_paragraph('Schemat procesu:')
add_code_block(
    'ETAP 1: Inicjalizacja Projektu → [ZATWIERDZENIE]\n'
    'ETAP 2: Analiza Stanu Obecnego (AS-IS) → [ZATWIERDZENIE]\n'
    'ETAP 3: Analiza Domeny Problemowej [DEEP] → [ZATWIERDZENIE]\n'
    'ETAP 4: Definicja Wymagań [DEEP] → [ZATWIERDZENIE]\n'
    'ETAP 5: Projekt Stanu Docelowego (TO-BE) → [ZATWIERDZENIE]\n'
    'ETAP 6: Analiza Luk i Harmonogram [DEEP] → [ZATWIERDZENIE]\n'
    'ETAP 7: Ocena Ryzyka → [ZATWIERDZENIE]\n'
    'ETAP 8: Biznesplan i ROI [DEEP] → [ZATWIERDZENIE]\n'
    'FINAŁ: Kompletny Pakiet Dokumentacji'
)

# --- STAGE 1 ---
doc.add_heading('4.1 Etap 1: Inicjalizacja Projektu i Mapa Interesariuszy', level=2)
doc.add_paragraph('Obszar wiedzy BABOK: Business Analysis Planning and Monitoring, Elicitation and Collaboration')
doc.add_paragraph('Tryb modelu: Standardowy')
doc.add_paragraph('Szacowany czas: 30–45 minut pracy + 1–2 dni na zebranie danych')

doc.add_heading('Cele etapu 1', level=3)
add_bullet('Wyjaśnienie zakresu i granic projektu')
add_bullet('Identyfikacja wszystkich interesariuszy i ich interesów')
add_bullet('Definicja kryteriów sukcesu (ilościowych i jakościowych)')
add_bullet('Ustanowienie planu komunikacji')
add_bullet('Dokumentacja założeń, zależności i ograniczeń')

doc.add_heading('Przebieg etapu 1', level=3)

doc.add_paragraph('Krok 1.1: Wyjaśnienie zakresu (Scope Clarification) — 5–6 pytań:')
add_bullet('O czym jest projekt? Opis problemu lub szansy biznesowej')
add_bullet('Jakie typy dokumentów/procesów są w zakresie?')
add_bullet('Jakie systemy IT są aktualnie używane? (ERP, księgowość, DMS)')
add_bullet('Jakie regulacje prawne mają znaczenie? (GDPR, e-fakturowanie, ISO)')
add_bullet('Jakie są terminy regulacyjne?')
add_bullet('Jaki jest szacunkowy budżet i harmonogram?')

doc.add_paragraph('Krok 1.2: Identyfikacja interesariuszy — 3 pytania:')
add_bullet('Kto jest sponsorem, decydentem budżetowym, liderem biznesowym/IT?')
add_bullet('Czy istnieją dodatkowe grupy interesariuszy?')
add_bullet('Czy są zaangażowani interesariusze zewnętrzni (audytorzy, konsultanci)?')

doc.add_paragraph('Krok 1.3: Definicja kryteriów sukcesu — 4 pytania:')
add_bullet('Jakie są obecne metryki bazowe (czas przetwarzania, koszty, błędy)?')
add_bullet('Jakie są realistyczne wartości docelowe?')
add_bullet('Jaki jest główny motywator projektu (#1)?')
add_bullet('Jakie są oczekiwania ROI (okres zwrotu, minimalna stopa zwrotu)?')

doc.add_heading('Dokument wynikowy etapu 1', level=3)
doc.add_paragraph('Plik: STAGE_01_Project_Initialization.md')
doc.add_paragraph('Zawiera:')
add_bullet('Streszczenie wykonawcze (1 strona)')
add_bullet('Zakres projektu (w zakresie / poza zakresem)')
add_bullet('Rejestr interesariuszy z analizą zainteresowania i wpływu')
add_bullet('Matrycę RACI (Responsible, Accountable, Consulted, Informed)')
add_bullet('Kryteria sukcesu z wartościami bazowymi i docelowymi')
add_bullet('Wymagania regulacyjne')
add_bullet('Plan komunikacji')
add_bullet('Ograniczenia projektu (budżet, harmonogram, zasoby)')
add_bullet('Założenia i zależności')
add_bullet('Otwarte pytania i decyzje do podjęcia')

# --- STAGE 2 ---
doc.add_heading('4.2 Etap 2: Analiza Stanu Obecnego (AS-IS)', level=2)
doc.add_paragraph('Obszar wiedzy BABOK: Strategy Analysis, Elicitation and Collaboration')
doc.add_paragraph('Tryb modelu: Standardowy (z opcjonalną głęboką analizą dla złożonych wąskich gardeł)')
doc.add_paragraph('Szacowany czas: 1–2 godziny pracy + 3–5 dni na zbieranie danych')

doc.add_heading('Cele etapu 2', level=3)
add_bullet('Dokumentacja obecnych procesów biznesowych (mapy procesów AS-IS)')
add_bullet('Identyfikacja punktów bólu (pain points) i nieefektywności')
add_bullet('Zebranie metryk bazowych (czas, koszt, wolumen, wskaźniki błędów)')
add_bullet('Mapowanie krajobazu systemowego i przepływów danych')
add_bullet('Ustalenie ilościowej bazy kosztowej do obliczenia ROI w etapie 8')

doc.add_heading('Przebieg etapu 2', level=3)

doc.add_paragraph('Krok 2.1: Mapowanie procesów AS-IS — agent pyta o:')
add_bullet('Główne procesy do zmapowania (top 3–5 procesów krytycznych)')
add_bullet('Opis bieżącego przepływu pracy: kto inicjuje, jakie kroki, jakie systemy')
add_bullet('Warianty procesów (np. krajowe vs. międzynarodowe)')
add_bullet('Najczęstsze wyjątki i scenariusze błędów')

doc.add_paragraph('Krok 2.2: Identyfikacja punktów bólu:')
add_bullet('Top 5 największych frustracji z obecnych procesów')
add_bullet('Obejścia (workarounds) stosowane przez pracowników')
add_bullet('Ograniczenia systemów IT')
add_bullet('Wpływ każdego problemu: czas stracony, częstotliwość, koszty')

doc.add_paragraph('Krok 2.3: Zbieranie metryk bazowych:')
add_bullet('Metryki wolumenu (liczba transakcji/miesiąc, liczba użytkowników)')
add_bullet('Metryki czasu (średni czas procesów, cykl zatwierdzania)')
add_bullet('Metryki kosztów (godziny FTE, koszty systemów, koszty fizyczne)')
add_bullet('Metryki jakości (wskaźnik błędów, incydenty compliance)')

doc.add_paragraph('Krok 2.4: Analiza systemów i przepływu danych:')
add_bullet('Jak dane przepływają między systemami? (ręcznie/API/plik/email)')
add_bullet('Problemy z jakością danych (duplikaty, niespójności)')
add_bullet('Dostęp i bezpieczeństwo danych')

doc.add_heading('Dokument wynikowy etapu 2', level=3)
doc.add_paragraph('Plik: STAGE_02_Current_State_Analysis.md')
doc.add_paragraph('Zawiera: mapy procesów AS-IS, rejestr punktów bólu, metryki bazowe, '
                  'analizę systemów, baza kosztowa, podsumowanie mocnych i słabych stron.')

# --- STAGE 3 ---
doc.add_heading('4.3 Etap 3: Analiza Domeny Problemowej', level=2)
doc.add_paragraph('Obszar wiedzy BABOK: Strategy Analysis, Requirements Analysis and Design Definition')
doc.add_paragraph('Tryb modelu: DEEP ANALYSIS MODE (automatyczny)')
doc.add_paragraph('Szacowany czas: 45–60 minut pracy + 1–2 dni na walidację')

doc.add_heading('Dlaczego tryb głębokiej analizy?', level=3)
add_bullet('Identyfikacja przyczyn źródłowych wymaga zaawansowanego rozumowania')
add_bullet('Priorytetyzacja obejmuje wieloczynnikowe podejmowanie decyzji')
add_bullet('Zależności między problemami mogą nie być oczywiste')

doc.add_heading('Techniki analityczne stosowane w etapie 3', level=3)

doc.add_paragraph('Analiza 5 Whys (5 Dlaczego):')
doc.add_paragraph(
    'Dla każdego krytycznego problemu agent wykonuje analizę „5 Dlaczego", czyli wielopoziomowe '
    'pytanie „dlaczego to się dzieje?" aż do dotarcia do przyczyny źródłowej (root cause). '
    'Dzięki temu rozwiązujemy przyczyny, a nie tylko objawy.'
)

doc.add_paragraph('Diagram Ishikawy (Fishbone):')
doc.add_paragraph(
    'Dla top 3 problemów agent tworzy diagramy przyczynowo-skutkowe analizujące czynniki w 6 obszarach: '
    'Metody (procesy), Maszyny (systemy), Materiały (dane), Ludzie (umiejętności), '
    'Pomiar (metryki), Otoczenie (regulacje).'
)

doc.add_paragraph('Macierz Impact-Effort (Wpływ-Wysiłek):')
doc.add_paragraph(
    'Wszystkie przyczyny źródłowe są umieszczane w macierzy 2×2: wysoki/niski wpływ '
    'vs. wysoki/niski wysiłek. Wynik: Quick Wins (szybkie zwycięstwa), Inicjatywy Strategiczne, '
    'Elementy do uzupełnienia, Elementy do odrzucenia.'
)

doc.add_heading('Dokument wynikowy etapu 3', level=3)
doc.add_paragraph('Plik: STAGE_03_Problem_Domain_Analysis.md')
doc.add_paragraph('Zawiera: kategoryzację problemów, analizę przyczyn źródłowych, diagramy Ishikawy, '
                  'macierz impact-effort, mapę zależności, zalecaną kolejność priorytetów.')

# --- STAGE 4 ---
doc.add_heading('4.4 Etap 4: Definicja Wymagań Rozwiązania', level=2)
doc.add_paragraph('Obszar wiedzy BABOK: Requirements Analysis and Design Definition, Requirements Life Cycle Management')
doc.add_paragraph('Tryb modelu: DEEP ANALYSIS MODE (automatyczny)')
doc.add_paragraph('Szacowany czas: 2–3 godziny pracy + 3–5 dni na walidację z interesariuszami')

doc.add_heading('Cele etapu 4', level=3)
add_bullet('Definicja wymagań funkcjonalnych (FR) adresujących przyczyny źródłowe z etapu 3')
add_bullet('Definicja wymagań niefunkcjonalnych (NFR): wydajność, bezpieczeństwo, dostępność')
add_bullet('Opracowanie user stories z kryteriami akceptacji')
add_bullet('Priorytetyzacja MoSCoW (Must/Should/Could/Won\'t)')
add_bullet('Budowa Macierzy Śledzenia Wymagań (RTM)')
add_bullet('Ustanowienie procesu kontroli zmian')

doc.add_heading('Priorytetyzacja MoSCoW', level=3)
add_table(
    ['Priorytet', 'Znaczenie', 'Opis'],
    [
        ['MUST', 'Obowiązkowe', 'Kluczowe dla uruchomienia. Projekt nie powiedzie się bez tego.'],
        ['SHOULD', 'Ważne', 'Istotne, ale projekt może startować bez tego (faza 1+).'],
        ['COULD', 'Pożądane', 'Miło mieć, ale nie krytyczne (backlog na przyszłość).'],
        ['WON\'T', 'Poza zakresem', 'Poza zakresem obecnego projektu.'],
    ]
)

doc.add_heading('Format user stories', level=3)
add_code_block(
    'US-001: Automatyczne rozpoznawanie faktur\n'
    'JAKO pracownik księgowości\n'
    'CHCĘ aby system automatycznie rozpoznawał dane z faktury\n'
    'PO TO ABY skrócić czas ręcznego wprowadzania danych\n\n'
    'Kryteria akceptacji:\n'
    '- AC-01: GIVEN faktura PDF WHEN przesłana do systemu THEN dane wyodrębnione w <30s\n'
    '- AC-02: GIVEN rozpoznane dane WHEN walidacja THEN dokładność >95%'
)

doc.add_heading('Dokument wynikowy etapu 4', level=3)
doc.add_paragraph('Plik: STAGE_04_Solution_Requirements.md')
doc.add_paragraph('Zawiera: pełną listę FR i NFR, user stories z kryteriami akceptacji, '
                  'macierz RTM, podsumowanie MoSCoW, proces kontroli zmian.')

# --- STAGE 5 ---
doc.add_heading('4.5 Etap 5: Projekt Stanu Docelowego (TO-BE)', level=2)
doc.add_paragraph('Obszar wiedzy BABOK: Requirements Analysis and Design Definition')
doc.add_paragraph('Tryb modelu: Standardowy (z głęboką analizą dla decyzji architektonicznych)')
doc.add_paragraph('Szacowany czas: 1–2 godziny pracy + 2–3 dni na walidację techniczną')

doc.add_heading('Cele etapu 5', level=3)
add_bullet('Zaprojektowanie architektury docelowej (adresującej wymagania MUST i SHOULD)')
add_bullet('Tworzenie map procesów TO-BE (zoptymalizowane przepływy)')
add_bullet('Definicja architektury integracji między systemami')
add_bullet('Projekt przepływu danych i modelu danych')
add_bullet('Porównanie projektu z wymaganiami (walidacja pokrycia)')

doc.add_heading('Kluczowe decyzje w etapie 5', level=3)
add_table(
    ['Decyzja', 'Opcje', 'Kryteria wyboru'],
    [
        ['Model wdrożenia', 'Chmura SaaS / IaaS / On-premise / Hybrydowy', 'Koszty, bezpieczeństwo, skalowalność'],
        ['Podejście', 'Kup gotowe / Zbuduj / Hybrydowe', 'Czas, koszty, dopasowanie do wymagań'],
        ['Integracja', 'API / Pliki / Zdarzenia', 'Dostępność API, wymagania czasu rzeczywistego'],
    ]
)

doc.add_heading('Dokument wynikowy etapu 5', level=3)
doc.add_paragraph('Plik: STAGE_05_Future_State_Design.md')
doc.add_paragraph('Zawiera: mapy procesów TO-BE, porównanie AS-IS vs. TO-BE, architekturę docelową, '
                  'architekturę integracji, architekturę danych, pokrycie wymagań.')

# --- STAGE 6 ---
doc.add_heading('4.6 Etap 6: Analiza Luk i Harmonogram Wdrożenia', level=2)
doc.add_paragraph('Obszar wiedzy BABOK: Strategy Analysis, Solution Evaluation')
doc.add_paragraph('Tryb modelu: DEEP ANALYSIS MODE (automatyczny)')
doc.add_paragraph('Szacowany czas: 1 godzina pracy + 1 dzień na walidację')

doc.add_heading('Cele etapu 6', level=3)
add_bullet('Systematyczne porównanie stanu AS-IS z projektem TO-BE (macierz luk)')
add_bullet('Definicja faz wdrożenia z jasnymi kamieniami milowymi')
add_bullet('Alokacja zasobów i szacowanie pracochłonności per faza')
add_bullet('Tworzenie planu zarządzania zmianą')
add_bullet('Definicja planu szkoleń')
add_bullet('Ustanowienie kryteriów uruchomienia produkcyjnego (go-live criteria)')

doc.add_heading('Typowe fazy wdrożenia', level=3)
add_table(
    ['Faza', 'Nazwa', 'Opis'],
    [
        ['Faza 0', 'Przygotowanie', 'Wybór dostawcy, konfiguracja środowisk, formowanie zespołu'],
        ['Faza 1', 'Quick Wins + Rdzeń', 'Podstawowa funkcjonalność + szybkie zwycięstwa z etapu 3'],
        ['Faza 2', 'Integracja + Automatyzacja', 'Integracje systemowe, automatyzacja procesów, migracja danych'],
        ['Faza 3', 'Optymalizacja + Zaawansowane', 'Wymagania SHOULD, zaawansowane funkcje, tuning wydajności'],
        ['Faza 4', 'Stabilizacja + Przekazanie', 'UAT, szkolenia, przygotowanie go-live, wsparcie hypercare'],
    ]
)

doc.add_heading('Dokument wynikowy etapu 6', level=3)
doc.add_paragraph('Plik: STAGE_06_Gap_Analysis_Roadmap.md')
doc.add_paragraph('Zawiera: macierz luk, fazy wdrożenia z kryteriami wejścia/wyjścia, '
                  'plan zasobów, plan zarządzania zmianą, plan szkoleń, kryteria go-live.')

# --- STAGE 7 ---
doc.add_heading('4.7 Etap 7: Ocena Ryzyka i Strategia Mitygacji', level=2)
doc.add_paragraph('Obszar wiedzy BABOK: Strategy Analysis, Business Analysis Planning and Monitoring')
doc.add_paragraph('Tryb modelu: Standardowy (z głęboką analizą dla DPIA i złożonych scenariuszy ryzyka)')
doc.add_paragraph('Szacowany czas: 45 minut pracy + 1 dzień na walidację')

doc.add_heading('Kategorie ryzyk', level=3)
add_table(
    ['Kategoria', 'Przykłady', 'Źródło identyfikacji'],
    [
        ['Techniczne', 'Złożoność integracji, ograniczenia API, jakość migracji danych', 'Etapy 5, 6'],
        ['Organizacyjne', 'Opór użytkowników, utrata kluczowego interesariusza, scope creep', 'Etapy 1, 4, 6'],
        ['Regulacyjne', 'Niedotrzymanie terminu, zmiana wymagań prawnych', 'Etap 1'],
        ['Finansowe', 'Przekroczenie budżetu, zmiana cen dostawcy', 'Etapy 1, 5, 6'],
        ['Zewnętrzne', 'Upadek dostawcy, zmiana warunków rynkowych', 'Etap 5'],
    ]
)

doc.add_heading('Strategie odpowiedzi na ryzyko', level=3)
add_bullet('Unikanie (Avoid) — zmiana podejścia aby wyeliminować ryzyko')
add_bullet('Transfer — przeniesienie ryzyka na stronę trzecią (dostawca, ubezpieczenie)')
add_bullet('Mitygacja (Mitigate) — redukcja prawdopodobieństwa lub wpływu')
add_bullet('Akceptacja (Accept) — uznanie ryzyka i przygotowanie planu awaryjnego')

doc.add_heading('DPIA — Ocena Skutków dla Ochrony Danych', level=3)
doc.add_paragraph(
    'Jeśli projekt obejmuje przetwarzanie danych osobowych, agent przygotuje ocenę DPIA '
    'zgodnie z art. 35 RODO. DPIA jest obowiązkowa gdy: wykorzystywane są nowe technologie, '
    'przetwarzanie na dużą skalę, systematyczny monitoring osób.'
)

doc.add_heading('Dokument wynikowy etapu 7', level=3)
doc.add_paragraph('Plik: STAGE_07_Risk_Assessment.md')
doc.add_paragraph('Zawiera: rejestr ryzyk, macierz priorytetyzacji, strategie mitygacji, '
                  'DPIA (jeśli wymagana), plan monitorowania ryzyk, rekomendacja rezerwy budżetowej.')

# --- STAGE 8 ---
doc.add_heading('4.8 Etap 8: Biznesplan i Model ROI', level=2)
doc.add_paragraph('Obszar wiedzy BABOK: Solution Evaluation, Strategy Analysis')
doc.add_paragraph('Tryb modelu: DEEP ANALYSIS MODE (automatyczny)')
doc.add_paragraph('Szacowany czas: 1–2 godziny pracy + 2–3 dni na walidację finansową')

doc.add_heading('Cele etapu 8', level=3)
add_bullet('Kalkulacja całkowitego kosztu posiadania (TCO — Total Cost of Ownership)')
add_bullet('Kwantyfikacja oczekiwanych korzyści (oszczędności, przychody, redukcja ryzyka)')
add_bullet('Budowa modelu finansowego z NPV, IRR i okresem zwrotu')
add_bullet('Analiza wrażliwości (scenariusze najlepszy/oczekiwany/najgorszy)')
add_bullet('Porównanie opcji inwestycyjnych (jeśli rozważano kilka podejść)')
add_bullet('Sformułowanie rekomendacji wykonawczej')

doc.add_heading('Kluczowe wskaźniki finansowe', level=3)
add_table(
    ['Wskaźnik', 'Opis', 'Typowy benchmark'],
    [
        ['NPV (Net Present Value)', 'Wartość bieżąca netto przepływów pieniężnych', '> 0'],
        ['IRR (Internal Rate of Return)', 'Wewnętrzna stopa zwrotu', '> WACC firmy (typowo 8–12%)'],
        ['Okres zwrotu (Payback Period)', 'Czas do odzyskania inwestycji', '< 18 miesięcy'],
        ['ROI 3-letni', 'Zwrot z inwestycji w horyzoncie 3 lat', '150–300% (branżowy benchmark)'],
        ['Benefit-Cost Ratio', 'Stosunek korzyści do kosztów', '> 1.5:1'],
    ]
)

doc.add_heading('Analiza wrażliwości', level=3)
doc.add_paragraph(
    'Agent tworzy 4 scenariusze: najlepszy (korzyści +20%, koszty -10%), oczekiwany '
    '(podstawowe założenia), najgorszy (korzyści -30%, koszty +20%, opóźnienie +3 mies.) '
    'oraz scenariusz opóźnienia. Ponadto wyznacza próg rentowności dla kluczowych parametrów.'
)

doc.add_heading('Dokument wynikowy etapu 8', level=3)
doc.add_paragraph('Plik: STAGE_08_Business_Case_ROI.md')
doc.add_paragraph('Zawiera: TCO, kwantyfikację korzyści, model finansowy (NPV/IRR/Payback), '
                  'analizę wrażliwości, porównanie opcji, rekomendację wykonawczą.')

doc.add_page_break()

# ==================== CHAPTER 5: COMMANDS ====================
doc.add_heading('5. Interfejs Poleceń — Pełna Referencja', level=1)

doc.add_paragraph(
    'Agent reaguje na polecenia w stylu terminala. Polecenia nie rozróżniają wielkości liter. '
    'Poniżej pełna lista dostępnych poleceń z opisami.'
)

doc.add_heading('5.1 Polecenia zarządzania projektem', level=2)
add_table(
    ['Polecenie', 'Opis', 'Przykład'],
    [
        ['BEGIN NEW PROJECT', 'Tworzy nowy projekt z unikalnym ID', 'BEGIN NEW PROJECT'],
        ['SAVE PROJECT', 'Zapisuje stan projektu (po zatwierdzeniu etapu)', 'SAVE PROJECT'],
        ['LOAD PROJECT [id]', 'Wznawia zapisany projekt', 'LOAD PROJECT BABOK-20260208-M3R1'],
        ['Pause', 'Wstrzymuje sesję (auto-zapis do dziennika)', 'Pause'],
        ['Status', 'Pokazuje postęp we wszystkich 8 etapach', 'Status'],
        ['Reset', 'Czyści dane i zaczyna od nowa (wymaga potwierdzenia)', 'Reset'],
        ['Version', 'Wyświetla wersję agenta', 'Version'],
    ]
)

doc.add_heading('5.2 Polecenia zarządzania etapami', level=2)
add_table(
    ['Polecenie', 'Opis', 'Przykład'],
    [
        ['Approve [N]', 'Zatwierdza etap N i przechodzi do następnego', 'Approve 1'],
        ['Reject [N] [powód]', 'Odrzuca etap z wyjaśnieniem', 'Reject 2 brak danych bazowych'],
        ['Skip to [N]', 'Przeskakuje do etapu (niezalecane, pokazuje ostrzeżenie)', 'Skip to 4'],
        ['Regenerate [N]', 'Odbudowuje etap od zera', 'Regenerate 3'],
        ['Validate [N]', 'Sprawdza kompletność przed zatwierdzeniem', 'Validate 1'],
    ]
)

doc.add_heading('5.3 Polecenia dokumentów', level=2)
add_table(
    ['Polecenie', 'Opis', 'Przykład'],
    [
        ['Export [N]', 'Eksportuje deliverable etapu N', 'Export 1'],
        ['Export all', 'Eksportuje wszystkie ukończone etapy', 'Export all'],
        ['Summary [N]', 'Pokazuje tylko streszczenie wykonawcze', 'Summary 2'],
        ['Detail [N]', 'Pokazuje pełną szczegółową analizę', 'Detail 2'],
        ['Preview [N]', 'Podgląd tego co zostanie wygenerowane', 'Preview 3'],
        ['Template [typ]', 'Pokazuje pusty szablon', 'Template stakeholder_register'],
        ['MAKE DOCX (CLI)', 'Generuje profesjonalne pliki DOCX', 'babok make docx BABOK-20260208-XXXX'],
        ['MAKE PDF (CLI)', 'Generuje pliki PDF', 'babok make pdf BABOK-20260208-XXXX'],
    ]
)

doc.add_heading('5.4 Polecenia analizy', level=2)
add_table(
    ['Polecenie', 'Opis', 'Tryb modelu', 'Przykład'],
    [
        ['Deep analysis [temat]', 'Aktywuje głęboką analizę AI', 'Premium', 'Deep analysis vendor_selection'],
        ['Quick check [zapytanie]', 'Szybkie zapytanie', 'Szybki', 'Quick check invoice_volume'],
        ['Compare [A] [B]', 'Porównuje dwie opcje', 'Premium', 'Compare cloud on-premise'],
        ['Calculate ROI [scenariusz]', 'Modelowanie finansowe', 'Premium', 'Calculate ROI full_automation'],
    ]
)

doc.add_heading('5.5 Polecenia współpracy i nawigacji pytań', level=2)
add_table(
    ['Polecenie', 'Opis'],
    [
        ['Show assumptions', 'Lista wszystkich bieżących założeń'],
        ['Show decisions', 'Lista wszystkich podjętych decyzji'],
        ['Show risks', 'Lista wszystkich zidentyfikowanych ryzyk'],
        ['Show requirements', 'Lista wszystkich wymagań (od etapu 4+)'],
        ['Update [item_id]', 'Modyfikuje konkretny element (np. Update FR-015)'],
        ['Next question', 'Przeskakuje do następnego pytania'],
        ['Previous question', 'Wraca do poprzedniego pytania'],
        ['Skip questions', 'Wyświetla wszystkie pozostałe pytania naraz (batch)'],
        ['Workshop [N]', 'Tryb interaktywny z częstym wkładem człowieka'],
        ['Async [N]', 'Tryb autonomiczny z minimalną interakcją'],
    ]
)

doc.add_page_break()

# ==================== CHAPTER 6: PROJECT STRUCTURE ====================
doc.add_heading('6. Struktura Projektu i Plików', level=1)

doc.add_heading('6.1 Zalecana struktura folderów', level=2)
doc.add_paragraph(
    'Poniższa struktura folderów jest rekomendowana dla każdego projektu BABOK. '
    'Wszystkie placeholders [w nawiasach] należy zastąpić rzeczywistymi nazwami.'
)

structure_items = [
    ('01_Project_Charter/', 'Karta projektu, rejestr interesariuszy, notatki z kickoffu'),
    ('02_Current_State/', 'Analiza stanu obecnego, mapy procesów AS-IS, metryki bazowe'),
    ('03_Problem_Analysis/', 'Diagramy Ishikawy, analiza przyczyn źródłowych, macierz impact-effort'),
    ('04_Requirements/', 'Wymagania, user stories, use cases, wireframy, macierz RTM'),
    ('05_Solution_Design/', 'Architektura TO-BE, diagramy, specyfikacje techniczne'),
    ('06_Implementation_Planning/', 'Harmonogram Gantta, plan zasobów, plan zmian, plan szkoleń'),
    ('07_Risk_Management/', 'Rejestr ryzyk, plany mitygacji, plany awaryjne'),
    ('08_Business_Case/', 'Biznesplan, model finansowy, oferty dostawców'),
    ('09_Vendor_Evaluation/', 'Dokument RFP, macierz porównawcza, notatki z demo (opcjonalnie)'),
    ('10_Governance/', 'Protokoły spotkań, dzienniki decyzji i zmian'),
    ('11_Compliance/', 'DPIA, dokumenty regulacyjne, opinie prawne (jeśli dotyczy)'),
    ('12_Testing/', 'Plan testów, przypadki testowe, wyniki UAT'),
    ('13_Training_Materials/', 'Podręczniki użytkownika, materiały szkoleniowe'),
    ('14_Communication/', 'Newsletter, FAQ, archiwum komunikacji'),
    ('FINAL_Documentation/', 'Skonsolidowana dokumentacja końcowa'),
]

add_table(
    ['Folder', 'Zawartość'],
    [[f, d] for f, d in structure_items]
)

doc.add_heading('6.2 Konwencja nazewnictwa plików', level=2)
doc.add_paragraph('Format:')
add_code_block('[NazwaDokumentu]_v[Major].[Minor]_[YYYY-MM-DD]_[Status].md')
doc.add_paragraph('Przykład:')
add_code_block(
    'STAGE_04_Solution_Requirements_v1.0_2026-02-15_DRAFT.md\n'
    'STAGE_04_Solution_Requirements_v1.1_2026-02-18_REVIEWED.md\n'
    'STAGE_04_Solution_Requirements_v2.0_2026-02-20_APPROVED.md'
)
doc.add_paragraph('Statusy:')
add_bullet('DRAFT — w trakcie prac')
add_bullet('REVIEWED — przejrzany przez interesariuszy')
add_bullet('APPROVED — formalnie zatwierdzony')
add_bullet('FINAL — wersja końcowa')

doc.add_heading('6.3 Wersjonowanie dokumentów', level=2)
doc.add_paragraph(
    'Każdy dokument powinien mieć numer wersji. Zmiany major (np. v1 → v2) oznaczają '
    'istotne przeróbki po recenzji. Zmiany minor (np. v1.0 → v1.1) oznaczają drobne poprawki.'
)

doc.add_heading('6.4 Bezpieczeństwo i kontrola dostępu', level=2)
add_table(
    ['Folder', 'Poziom dostępu', 'Uzasadnienie'],
    [
        ['01_Project_Charter', 'Sponsor, BA, PM', 'Dokument karty projektu'],
        ['04_Requirements', 'BA, Dev Team, Eksperci', 'Specyfikacje techniczne'],
        ['08_Business_Case', 'CFO, Sponsor, BA', 'Dane finansowe — ograniczony dostęp'],
        ['11_Compliance', 'Prawnik, Compliance, BA', 'Dokumenty prawne — ograniczony dostęp'],
    ]
)

add_note('Zasada: jeśli folder zawiera dane finansowe, prawne lub wrażliwe dane osobowe → Ograniczony Dostęp')

doc.add_page_break()

# ==================== CHAPTER 7: AGENT MODES ====================
doc.add_heading('7. Tryby Pracy Agenta', level=1)

doc.add_paragraph(
    'BABOK ANALYST automatycznie dobiera poziom zaawansowania analizy AI w zależności od '
    'złożoności zadania. Użytkownik może też ręcznie aktywować konkretny tryb.'
)

doc.add_heading('7.1 Tryb głębokiej analizy (Deep Analysis Mode)', level=2)
doc.add_paragraph('Automatycznie aktywowany w etapach: 3, 4, 6, 8')
doc.add_paragraph('Wyzwalacze ręczne:')
add_bullet('Polecenie: Deep analysis [temat]')
add_bullet('Decyzje krytyczne oznaczone przez użytkownika')
add_bullet('Konflikty między interesariuszami')
add_bullet('Nowe problemy nieobjęte standardowym frameworkiem BABOK')
doc.add_paragraph('Wskaźnik w rozmowie:')
add_code_block('[DEEP ANALYSIS MODE ACTIVATED]\nModel: Gemini Pro 3 / Claude Opus 4.6\nReasoning: [uzasadnienie]')

doc.add_heading('7.2 Tryb standardowy (Standard Mode)', level=2)
doc.add_paragraph(
    'Domyślny tryb pracy. Używany do większości zadań analitycznych: dokumentacja wymagań, '
    'wywiady z interesariuszami, mapowanie procesów, standardowe procedury BABOK.'
)

doc.add_heading('7.3 Tryb szybki (Rapid Mode)', level=2)
doc.add_paragraph('Wyzwalacze:')
add_bullet('Formatowanie dokumentów')
add_bullet('Pobieranie informacji z poprzednich etapów')
add_bullet('Proste pytania wyjaśniające')
add_bullet('Wypełnianie szablonów')
add_bullet('Sprawdzanie list kontrolnych')

doc.add_page_break()

# ==================== CHAPTER 8: JOURNAL ====================
doc.add_heading('8. Dziennik Projektu (Project Journal)', level=1)

doc.add_paragraph(
    'Każdy projekt BABOK utrzymuje plik dziennika (journal) w formacie JSON, który rejestruje '
    'każdą zmianę stanu. Dzięki temu możliwe jest wznowienie projektu dokładnie w miejscu '
    'przerwania.'
)

doc.add_heading('Format identyfikatora projektu', level=2)
add_code_block('BABOK-YYYYMMDD-XXXX\n\nYYYYMMDD — data utworzenia projektu\nXXXX — 4-znakowy losowy sufiks alfanumeryczny')

doc.add_heading('Zdarzenia rejestrowane w dzienniku', level=2)
add_table(
    ['Zdarzenie', 'Wyzwalacz', 'Rejestrowane dane'],
    [
        ['project_created', 'BEGIN NEW PROJECT', 'ID projektu, nazwa, znacznik czasu'],
        ['stage_started', 'Wejście w nowy etap', 'Numer etapu, znacznik czasu'],
        ['stage_completed', 'Wygenerowanie deliverable', 'Numer etapu, plik deliverable'],
        ['stage_approved', 'Approve [N]', 'Numer etapu, znacznik czasu, zatwierdzający'],
        ['stage_rejected', 'Reject [N]', 'Numer etapu, powód odrzucenia'],
        ['project_saved', 'SAVE PROJECT', 'Pełna migawka stanu'],
        ['project_loaded', 'LOAD PROJECT', 'ID projektu, etap wznowienia'],
        ['decision_made', 'Decyzja kluczowa', 'ID decyzji, opis, uzasadnienie'],
        ['assumption_added', 'Założenie agenta', 'ID założenia, opis, pewność'],
    ]
)

doc.add_page_break()

# ==================== CHAPTER 9: BEST PRACTICES ====================
doc.add_heading('9. Najlepsze Praktyki i Wskazówki', level=1)

doc.add_heading('Zalecenia (DO):', level=2)
add_bullet('Bądź konkretny — „średnio 50 faktur/miesiąc" jest lepsze niż „sporo"')
add_bullet('Czytaj sekcje streszczenia wykonawczego — szybki przegląd kluczowych wniosków')
add_bullet('Zatwierdzaj progresywnie — nie czekaj na perfekcję, iteruj')
add_bullet('Koryguj natychmiast — jeśli agent popełni błąd, popraw go od razu (CORRECTION:)')
add_bullet('Pytaj o wyjaśnienia — jeśli coś jest niejasne: „WHY did you conclude X?"')
add_bullet('Używaj Summary przed Approve — zawsze przejrzyj streszczenie przed zatwierdzeniem')
add_bullet('Regularnie eksportuj — nie czekaj do końca, używaj Export all po każdym zatwierdzeniu')
add_bullet('Używaj Workshop dla etapu 4 (wymagania) — wymaga dużo interakcji z interesariuszami')
add_bullet('Używaj Async dla etapu 2 (stan obecny) — agent może analizować dane samodzielnie')

doc.add_heading('Przestrogi (DON\'T):', level=2)
add_bullet('Nie zgaduj — jeśli nie wiesz, powiedz „NIE WIEM"')
add_bullet('Nie pomijaj pytań — każde pytanie agenta ma konkretny powód')
add_bullet('Nie zatwierdzaj na ślepo — przeczytaj dokument przed zatwierdzeniem')
add_bullet('Nie mieszaj etapów — dokończ jeden etap przed przejściem do następnego')
add_bullet('Nie używaj Deep Analysis do prostych pytań — zarezerwuj na krytyczne decyzje')

doc.add_page_break()

# ==================== CHAPTER 10: FAQ ====================
doc.add_heading('10. Rozwiązywanie Problemów (FAQ)', level=1)

problems = [
    ('Agent nie rozumie mojej odpowiedzi',
     'Przeformułuj bardziej szczegółowo. Zamiast „mamy dużo faktur" napisz '
     '„otrzymujemy średnio 450 faktur zakupowych miesięcznie".'),
    ('Agent pyta o coś, czego nie możemy określić teraz',
     'Napisz: „To pytanie wymaga decyzji zarządu. Zaznacz jako OPEN QUESTION i przejdźmy dalej." '
     'Agent oznaczy to jako otwarte i wróci do tematu w późniejszym etapie.'),
    ('Pomyliłem się w odpowiedzi kilka pytań temu',
     'Użyj CORRECTION w [Section X.Y]: [opis błędu i poprawka]. Agent cofnie się i poprawi.'),
    ('Agent wygenerował zbyt techniczny dokument',
     'Napisz: „Simplify section [X] for non-technical audience." Agent uprości język.'),
    ('Chcę zmienić priorytet wymagania',
     'Napisz: „In Stage 4, change requirement FR-015 from MUST to SHOULD" z uzasadnieniem.'),
    ('Nie pamiętam na jakim etapie jestem',
     'Użyj polecenia Status — agent pokaże postęp we wszystkich 8 etapach.'),
    ('Chcę zmienić coś po zatwierdzeniu etapu',
     'Użyj Regenerate [N] aby odbudować etap, potem Update [sekcja] i ponowne Approve [N].'),
    ('Polecenie nie działa',
     'Sprawdź pisownię. Polecenia nie rozróżniają wielkości liter, '
     'ale muszą być wpisane poprawnie.'),
]

for problem, solution in problems:
    doc.add_heading(f'Problem: {problem}', level=3)
    doc.add_paragraph(f'Rozwiązanie: {solution}')

doc.add_page_break()

# ==================== CHAPTER 11: SECURITY ====================
doc.add_heading('11. Bezpieczeństwo i Prywatność', level=1)

doc.add_heading('Dane, które agent NIE przechowuje:', level=2)
add_bullet('Haseł ani kluczy API')
add_bullet('Numerów kont bankowych klientów/dostawców')
add_bullet('Szczegółowych danych osobowych poza niezbędnymi do analizy')

doc.add_heading('Dane, które agent MOŻE przechowywać:', level=2)
add_bullet('Strukturę organizacyjną (role, nie nazwiska — chyba że podane)')
add_bullet('Procesy biznesowe (przepływy pracy)')
add_bullet('Zagregowane metryki finansowe (nie szczegółowe transakcje)')
add_bullet('Nazwy systemów i narzędzi')

doc.add_heading('Zgodność z regulacjami:', level=2)
add_bullet('GDPR — minimalizacja danych, prawo do usunięcia')
add_bullet('BABOK® Code of Conduct')
add_bullet('Zasady ISO 27001 (dokumentacja procesów)')

doc.add_paragraph(
    'Po zakończeniu projektu możesz usunąć konwersację z platformy AI, '
    'a pliki .md zachować lokalnie.'
)

doc.add_page_break()

# ==================== CHAPTER 12: TIME ESTIMATES ====================
doc.add_heading('12. Szacunkowy Czas Realizacji', level=1)

add_table(
    ['Etap', 'Czas pracy z agentem', 'Czas na konsultacje', 'Łącznie'],
    [
        ['Etap 1: Inicjalizacja', '30–45 min', '1–2 dni', '1–2 dni'],
        ['Etap 2: Stan obecny (AS-IS)', '1–2 godz.', '3–5 dni', '1 tydzień'],
        ['Etap 3: Analiza problemów', '45–60 min', '1–2 dni', '2–3 dni'],
        ['Etap 4: Wymagania', '2–3 godz.', '3–5 dni', '1 tydzień'],
        ['Etap 5: Stan docelowy (TO-BE)', '1–2 godz.', '2–3 dni', '3–4 dni'],
        ['Etap 6: Analiza luk', '1 godz.', '1 dzień', '1–2 dni'],
        ['Etap 7: Ocena ryzyka', '45 min', '1 dzień', '1–2 dni'],
        ['Etap 8: Biznesplan i ROI', '1–2 godz.', '2–3 dni', '3–5 dni'],
        ['RAZEM', '8–12 godz.', '2–3 tygodnie', '3–4 tygodnie'],
    ]
)

add_note('Większość czasu to NIE praca z agentem, lecz zbieranie danych od '
         'interesariuszy i konsultacje wewnętrzne.')

doc.add_page_break()

# ==================== CHAPTER 13: GLOSSARY ====================
doc.add_heading('13. Słowniczek Pojęć', level=1)

glossary_items = [
    ('BABOK®', 'Business Analysis Body of Knowledge — międzynarodowy standard analizy biznesowej (IIBA)'),
    ('AS-IS', 'Stan obecny procesów, systemów i organizacji (przed zmianą)'),
    ('TO-BE', 'Stan docelowy procesów, systemów i organizacji (po wdrożeniu)'),
    ('Deliverable', 'Dokument wynikowy etapu — formalny produkt pracy'),
    ('Deep Analysis Mode', 'Tryb zaawansowanej analizy AI dla decyzji krytycznych'),
    ('DPIA', 'Data Protection Impact Assessment — ocena skutków przetwarzania danych (RODO art. 35)'),
    ('FR', 'Functional Requirement — wymaganie funkcjonalne (co system ma robić)'),
    ('Gap Analysis', 'Analiza luk — porównanie stanu obecnego z docelowym'),
    ('IRR', 'Internal Rate of Return — wewnętrzna stopa zwrotu'),
    ('Ishikawa Diagram', 'Diagram przyczynowo-skutkowy (diagram rybiej ości)'),
    ('Journal', 'Dziennik projektu — plik JSON śledzący wszystkie zmiany stanu'),
    ('KPI', 'Key Performance Indicator — kluczowy wskaźnik efektywności'),
    ('MoSCoW', 'Must/Should/Could/Won\'t — metoda priorytetyzacji wymagań'),
    ('NFR', 'Non-Functional Requirement — wymaganie niefunkcjonalne (jak dobrze system ma działać)'),
    ('NPV', 'Net Present Value — wartość bieżąca netto'),
    ('Pain Point', 'Punkt bólu — problem lub frustracja w obecnym procesie'),
    ('RACI', 'Responsible/Accountable/Consulted/Informed — macierz odpowiedzialności'),
    ('ROI', 'Return on Investment — zwrot z inwestycji'),
    ('Root Cause', 'Przyczyna źródłowa — fundamentalny powód problemu'),
    ('RTM', 'Requirements Traceability Matrix — macierz śledzenia wymagań'),
    ('Scope', 'Zakres projektu — granice tego co jest w projekcie i co nie'),
    ('Stakeholder', 'Interesariusz — osoba lub grupa mająca interes w projekcie'),
    ('TCO', 'Total Cost of Ownership — całkowity koszt posiadania'),
    ('UAT', 'User Acceptance Testing — testy akceptacyjne użytkowników'),
    ('User Story', 'Historia użytkownika — opis funkcjonalności z perspektywy użytkownika'),
    ('WACC', 'Weighted Average Cost of Capital — średni ważony koszt kapitału'),
    ('5 Whys', '5 Dlaczego — technika dotarcia do przyczyny źródłowej przez wielokrotne pytanie „dlaczego?"'),
]

add_table(
    ['Pojęcie', 'Definicja'],
    [[term, defn] for term, defn in glossary_items]
)

# ==================== FOOTER ====================
doc.add_page_break()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('BABOK ANALYST — Podręcznik Użytkownika v1.8.1')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer2.add_run('© 2026 BABOK Agent Development Team')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

footer3 = doc.add_paragraph()
footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer3.add_run('Framework: BABOK® v3 (International Institute of Business Analysis)')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ==================== SAVE ====================
output_path = os.path.join(os.path.dirname(__file__), 'BABOK ANALYST - podrecznik uzytkownika.docx')
doc.save(output_path)
print(f"✅ Podręcznik zapisany: {output_path}")
